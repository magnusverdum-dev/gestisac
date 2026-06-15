from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\josefeio\Desktop\git\gestisac")
SOURCE_HTML = ROOT / "apps" / "web" / "public" / "gcsoftware-funcionalidades.html"
OUTPUT_DOCX = ROOT / "deliverables" / "GCSoftware_levantamento_funcional_assinatura.docx"


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "C7CFD9")
        borders.append(elem)
    tblPr.append(borders)


def set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_repeat_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_keep_together(paragraph) -> None:
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True


def add_run_text(paragraph, text: str, bold: bool = False, color: str | None = None, size: int = 11) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def extract_inventory() -> list[dict]:
    html = SOURCE_HTML.read_text(encoding="utf-8")
    start = html.find("const inventory = [")
    end = html.find("\n      const state =", start)
    if start == -1 or end == -1:
        raise RuntimeError("Não foi possível localizar o inventário no HTML fonte.")
    array_source = html[start + len("const inventory = ") : end].strip()
    if array_source.endswith(";"):
        array_source = array_source[:-1].rstrip()

    node_script = r"""
const fs = require('fs');
const src = fs.readFileSync(0, 'utf8');
const inventory = Function('"use strict"; return (' + src + ');')();
process.stdout.write(JSON.stringify(inventory));
"""
    proc = subprocess.run(
        [
            r"C:\Users\josefeio\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe",
            "-e",
            node_script,
        ],
        input=array_source.encode("utf-8"),
        text=False,
        capture_output=True,
        check=True,
    )
    data = json.loads(proc.stdout.decode("utf-8"))
    return data


def build_doc() -> Document:
    inventory = extract_inventory()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in [
        ("Heading 1", 15, "0F2747", 12, 6),
        ("Heading 2", 12, "0F2747", 10, 4),
        ("Heading 3", 11, "1F4D78", 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    add_run_text(title, "GCSoftware - levantamento funcional para assinatura", bold=True, color="000000", size=18)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    add_run_text(
        subtitle,
        "Documento simples para revisão manual do MVP. Marque as funcionalidades pretendidas e use a área final para assinatura.",
        color="4B5563",
        size=10,
    )

    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_borders(meta)
    meta.columns[0].width = Inches(2.2)
    meta.columns[1].width = Inches(2.2)
    meta.columns[2].width = Inches(2.1)
    meta_cells = meta.rows[0].cells
    meta_items = [
        ("Projecto", "GCSoftware / GESTISAC"),
        ("Formato", "PDF para impressão"),
        ("Uso", "Selecção e assinatura manual"),
    ]
    for cell, (label, value) in zip(meta_cells, meta_items):
        set_cell_width(cell, 2.2 if label != "Uso" else 2.1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_run_text(p, f"{label}\n", bold=True, color="0F2747", size=9)
        add_run_text(p, value, color="111827", size=10)

    doc.add_paragraph("")

    legend = doc.add_paragraph()
    legend.paragraph_format.space_after = Pt(6)
    legend.paragraph_format.space_before = Pt(0)
    add_run_text(legend, "Legenda: ", bold=True, color="0F2747", size=9)
    add_run_text(legend, "□ ", color="0F2747", size=9)
    add_run_text(legend, "assinalar manualmente  ", color="4B5563", size=9)
    add_run_text(legend, "| ", color="9CA3AF", size=9)
    add_run_text(legend, "verificado", bold=True, color="166534", size=9)
    add_run_text(legend, "  |  ", color="9CA3AF", size=9)
    add_run_text(legend, "docs", bold=True, color="0F2747", size=9)
    add_run_text(legend, "  |  ", color="9CA3AF", size=9)
    add_run_text(legend, "a confirmar", bold=True, color="9A6700", size=9)

    total_count = 0
    selected_count = 0

    for module_index, module in enumerate(inventory):
        if module_index > 0:
            doc.add_page_break()

        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.space_after = Pt(3)
        add_run_text(heading, f"{module['code']}  {module['title']}", bold=True, color="0F2747", size=15)
        summary = doc.add_paragraph()
        summary.paragraph_format.space_after = Pt(8)
        add_run_text(summary, module["summary"], color="4B5563", size=10)

        for group in module["groups"]:
            group_heading = doc.add_paragraph(style="Heading 2")
            group_heading.paragraph_format.space_after = Pt(2)
            add_run_text(group_heading, group["title"], bold=True, color="1F4D78", size=11)

            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            set_table_borders(table)
            widths = [0.42, 4.95, 1.13]
            for idx, width in enumerate(widths):
                table.columns[idx].width = Inches(width)

            hdr = table.rows[0].cells
            headers = ["", "Funcionalidade", "Estado"]
            for idx, text in enumerate(headers):
                hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_width(hdr[idx], widths[idx])
                set_cell_shading(hdr[idx], "E8EEF5")
                p = hdr[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                add_run_text(p, text, bold=True, color="0F2747", size=9)
            set_repeat_header(table.rows[0])

            for item in group["items"]:
                total_count += 1
                if item["status"] == "verificado":
                    selected_count += 1

                row = table.add_row()
                cells = row.cells
                cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_width(cells[0], widths[0])
                set_cell_width(cells[1], widths[1])
                set_cell_width(cells[2], widths[2])

                p0 = cells[0].paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p0.paragraph_format.space_after = Pt(0)
                add_run_text(p0, "☐", bold=True, color="111827", size=12)

                p1 = cells[1].paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p1.paragraph_format.space_after = Pt(0)
                add_run_text(p1, item["label"], bold=False, color="111827", size=10)
                if item.get("detail"):
                    p1.add_run("\n")
                    add_run_text(p1, item["detail"], bold=False, color="6B7280", size=9)

                p2 = cells[2].paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.paragraph_format.space_after = Pt(0)
                status = item["status"]
                status_color = {
                    "verificado": "166534",
                    "docs": "0F2747",
                    "pendente": "9A6700",
                    "a-confirmar": "9B1C1C",
                }.get(status, "4B5563")
                add_run_text(p2, status, bold=True, color=status_color, size=9)

            doc.add_paragraph("")

    doc.add_page_break()
    sig_title = doc.add_paragraph(style="Heading 1")
    add_run_text(sig_title, "Assinatura final", bold=True, color="0F2747", size=15)

    sig_text = doc.add_paragraph()
    sig_text.paragraph_format.space_after = Pt(8)
    add_run_text(
        sig_text,
        "Após a revisão, assinale a aprovação da lista de funcionalidades para avançar com o levantamento final e o desenvolvimento.",
        color="4B5563",
        size=10,
    )

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    set_table_borders(sig_table)
    sig_widths = [3.2, 3.3]
    for idx, width in enumerate(sig_widths):
        sig_table.columns[idx].width = Inches(width)
    sig_table.rows[0].cells[0].text = ""
    sig_table.rows[0].cells[1].text = ""
    sig_table.rows[1].cells[0].text = ""
    sig_table.rows[1].cells[1].text = ""

    labels = [
        ("Assinatura da cliente", "Nome: _________________________________\nAssinatura: ____________________________"),
        ("Data e observações", "Data: __________________  Hora: __________________\nObservações: ____________________________"),
    ]
    for row_idx, (label, body) in enumerate(labels):
        row = sig_table.rows[row_idx]
        for cell_idx, text in enumerate([label, body]):
            cell = row.cells[cell_idx]
            set_cell_width(cell, sig_widths[cell_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_text(p, text, bold=(cell_idx == 0), color="111827", size=10 if cell_idx == 0 else 11)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(10)
    footer.paragraph_format.space_after = Pt(0)
    add_run_text(
        footer,
        f"Total de funcionalidades registadas: {total_count} | Verificadas no demo: {sum(1 for module in inventory for group in module['groups'] for item in group['items'] if item['status'] == 'verificado')}",
        color="4B5563",
        size=9,
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return doc


if __name__ == "__main__":
    build_doc()
    print(OUTPUT_DOCX)
