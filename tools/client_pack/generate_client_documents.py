from __future__ import annotations

import textwrap
from html import escape
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "apresentacao-cliente-gestisac"
ASSETS = OUT / "assets"
TODAY = date.today().strftime("%d/%m/%Y")

BLUE = "1F4D78"
MID_BLUE = "2E74B5"
DARK = "0B2545"
MUTED = "667085"
LIGHT = "F2F4F7"
SOFT_BLUE = "E8EEF5"
GOLD = "D7A84F"
GREEN = "2E7D32"
AMBER = "B36B00"
RED = "9B1C1C"


MODULES = [
    {
        "area": "Hoje / Dashboard",
        "purpose": "Centro de comando diario para acompanhar prioridades, indicadores e proximas acoes.",
        "users": "HQ/Admin, funcionario e cliente com conteudo filtrado por perfil.",
        "state": "Implementado",
        "next": "Afinar indicadores finais com regras internas da empresa e thresholds de prioridade.",
    },
    {
        "area": "Condominios",
        "purpose": "Ficha operacional de cada condominio, com estrutura, fraccoes, equipamentos, pedidos e relacoes.",
        "users": "HQ/Admin e perfis autorizados.",
        "state": "Parcial",
        "next": "Completar regras de dados obrigatorios, documentos associados e relatorios por condominio.",
    },
    {
        "area": "Equipa",
        "purpose": "Visao real de utilizadores/funcionarios, carga operacional, validacoes e atividade recente.",
        "users": "HQ/Admin; funcionarios veem apenas o necessario para execucao.",
        "state": "Implementado",
        "next": "Validar permissoes, categorias de equipa e metricas de produtividade que a empresa quer acompanhar.",
    },
    {
        "area": "Tarefas",
        "purpose": "Vista composta do trabalho diario a partir de tickets, manutencao, vistorias e agenda.",
        "users": "HQ/Admin e funcionarios.",
        "state": "Implementado",
        "next": "Formalizar regras de prioridade, SLA, validacao e escalonamento.",
    },
    {
        "area": "Pedidos / Tickets",
        "purpose": "Registo, triagem, acompanhamento, atribuicao, execucao e historico de pedidos.",
        "users": "HQ/Admin, funcionarios e clientes com niveis diferentes de acesso.",
        "state": "Implementado",
        "next": "Validar estados finais, templates de comunicacao e regras de notificacao.",
    },
    {
        "area": "Agenda",
        "purpose": "Planeamento de eventos, visitas, prazos e atividades ligadas a pedidos e condominios.",
        "users": "HQ/Admin, funcionarios e clientes quando aplicavel.",
        "state": "Implementado",
        "next": "Validar sincronizacao externa, lembretes e tipos oficiais de eventos.",
    },
    {
        "area": "Documentos",
        "purpose": "Organizar documentos por contexto: condominio, pedido, utilizador ou area administrativa.",
        "users": "HQ/Admin, clientes e funcionarios conforme permissao.",
        "state": "Parcial",
        "next": "Definir estrutura documental, permissoes, modelos e retencao.",
    },
    {
        "area": "Contabilidade",
        "purpose": "Apoiar quotas, despesas, saldos, avisos, recibos e controlo financeiro por condominio.",
        "users": "HQ/Admin e perfis financeiros.",
        "state": "Planeado",
        "next": "Validar processos atuais, mapas obrigatorios, integracoes e calendario de cobranca.",
    },
    {
        "area": "Relatorios",
        "purpose": "Criar visoes executivas e operacionais para gestao, clientes e auditoria.",
        "users": "HQ/Admin; clientes quando houver relatorios partilhados.",
        "state": "Planeado",
        "next": "Definir formatos, periodicidade, destinatarios e indicadores oficiais.",
    },
    {
        "area": "Manutencao / Vistorias",
        "purpose": "Registar intervencoes, inspecoes, estado, responsavel e ligacao a pedidos/tarefas.",
        "users": "HQ/Admin e funcionarios.",
        "state": "Parcial",
        "next": "Validar checklists, evidencias fotograficas, assinaturas e validacao final.",
    },
    {
        "area": "Fornecedores",
        "purpose": "Centralizar contactos, especialidades, historico e ligacao a intervencoes.",
        "users": "HQ/Admin.",
        "state": "Planeado",
        "next": "Definir categorias, avaliacao, documentos obrigatorios e regras de adjudicacao.",
    },
]


REQUIREMENTS = [
    ("RF-001", "Acesso", "Entrada em ambiente de demonstracao sem escrita manual de credenciais.", "Alta", "Implementado", "Contrato loginless para desenvolvimento e smoke tests.", ""),
    ("RF-002", "Acesso", "Separacao de perfis HQ/Admin, Funcionario e Cliente.", "Alta", "Implementado", "Cada contexto tem menu e permissoes filtradas.", ""),
    ("RF-003", "Acesso", "Gestao de utilizadores e equipa com dados reais da base de dados.", "Alta", "Implementado", "Endpoint /api/team expõe apenas dados publicos e metricas operacionais.", ""),
    ("RF-004", "Hoje", "Dashboard operacional com indicadores e prioridades do dia.", "Alta", "Implementado", "Centro de comando para apresentacao e operacao diaria.", ""),
    ("RF-005", "Hoje", "Cartoes/resumos de modulos principais para entrada rapida.", "Media", "Implementado", "A confirmar se estes sao os quatro cartoes finais para cliente.", ""),
    ("RF-006", "Condominios", "Listar condominios com estado, localidade, fraccoes e completude.", "Alta", "Implementado", "Vista operacional densa, sem cartoes repetitivos.", ""),
    ("RF-007", "Condominios", "Abrir ficha operacional do condominio e ver relacoes ligadas.", "Alta", "Parcial", "Dados base presentes; regras finais de ficha a validar.", ""),
    ("RF-008", "Condominios", "Gerir fraccoes, blocos, elevadores e equipamentos.", "Alta", "Parcial", "Funcionalidade ligada ao modelo; validar campos obrigatorios.", ""),
    ("RF-009", "Condominios", "Associar documentos e historico ao condominio.", "Media", "Parcial", "Depende de estrutura documental final.", ""),
    ("RF-010", "Equipa", "Listar membros da equipa com papel, contacto e carga operacional.", "Alta", "Implementado", "Usa users e atribuicoes existentes.", ""),
    ("RF-011", "Equipa", "Ver tarefas/pedidos atribuidos por funcionario.", "Alta", "Implementado", "Ligado a tickets.assigned_worker_id e inspections.assigned_worker_id.", ""),
    ("RF-012", "Equipa", "Validar perfis, departamentos e permissoes detalhadas.", "Alta", "A validar", "Precisa de matriz de autoridade real da empresa.", ""),
    ("RF-013", "Tarefas", "Agregacao operacional de tickets, manutencao, vistorias e eventos.", "Alta", "Implementado", "Nao cria tabela tasks; e uma vista composta.", ""),
    ("RF-014", "Tarefas", "Filtros de hoje, em curso, validacao e atrasadas.", "Alta", "Implementado", "Ajustar criterios com a empresa.", ""),
    ("RF-015", "Tarefas", "Acao seguinte e responsavel visiveis no painel contextual.", "Alta", "Implementado", "Padrao visual transversal.", ""),
    ("RF-016", "Pedidos", "Criar, listar, abrir e acompanhar pedidos/tickets.", "Alta", "Implementado", "Ciclo demonstravel por perfil.", ""),
    ("RF-017", "Pedidos", "Atribuir pedido a funcionario.", "Alta", "Implementado", "Usa assigned_worker_id.", ""),
    ("RF-018", "Pedidos", "Atualizar estado do pedido e acompanhar historico.", "Alta", "Implementado", "Validar nomes finais dos estados.", ""),
    ("RF-019", "Pedidos", "Permitir validacao final pela administracao.", "Alta", "Parcial", "Fluxo existe conceptualmente; regras finais a confirmar.", ""),
    ("RF-020", "Pedidos", "Anexos, fotos e documentos no pedido.", "Media", "Parcial", "Estrutura existe/planeada; validar armazenamento e limites.", ""),
    ("RF-021", "Agenda", "Ver eventos e proximas atividades por perfil.", "Alta", "Implementado", "Ligacao a calendar_events.", ""),
    ("RF-022", "Agenda", "Criar e relacionar eventos com pedido, tarefa ou condominio.", "Alta", "Parcial", "Confirmar regras de criacao e notificacao.", ""),
    ("RF-023", "Agenda", "Lembretes e notificacoes automáticas.", "Media", "Planeado", "Definir canais: email, push, SMS ou WhatsApp.", ""),
    ("RF-024", "Documentos", "Biblioteca documental por condominio/pedido/utilizador.", "Alta", "Parcial", "Definir taxonomia documental.", ""),
    ("RF-025", "Documentos", "Permissoes de visualizacao por perfil.", "Alta", "A validar", "Depende de regras legais e internas.", ""),
    ("RF-026", "Contabilidade", "Registo de quotas, pagamentos e dividas.", "Alta", "Planeado", "Validar processo financeiro atual antes de implementar.", ""),
    ("RF-027", "Contabilidade", "Mapas financeiros e extratos para clientes.", "Alta", "Planeado", "Formato deve ser aprovado pela administracao.", ""),
    ("RF-028", "Relatorios", "Relatorios de operacao, SLA, pedidos e manutencao.", "Media", "Planeado", "Definir periodicidade e destinatarios.", ""),
    ("RF-029", "Vistorias", "Planeamento, execucao e validacao de vistorias.", "Alta", "Parcial", "Ligado a inspections e tarefas.", ""),
    ("RF-030", "Fornecedores", "Cadastro e historico de fornecedores.", "Media", "Planeado", "Necessita validacao de campos legais/operacionais.", ""),
    ("RNF-001", "Seguranca", "API protegida, sem expor segredos no frontend.", "Alta", "Implementado", "Rust/Axum e variaveis de ambiente controladas.", ""),
    ("RNF-002", "Seguranca", "Separacao de dados por tenant/cliente.", "Alta", "Implementado", "Modelo multi-tenant existente.", ""),
    ("RNF-003", "Auditoria", "Historico de alteracoes relevantes.", "Alta", "Parcial", "Definir eventos obrigatorios e prazo de retencao.", ""),
    ("RNF-004", "Performance", "Interface rapida e instalavel como PWA.", "Alta", "Implementado", "Qwik e PWA para mobile/desktop.", ""),
    ("RNF-005", "Disponibilidade", "Deploy publicado com checks, smoke tests e warmup.", "Alta", "Implementado", "Vercel/GitHub e scripts de producao.", ""),
    ("RNF-006", "Backups", "Backups e recuperacao de base de dados.", "Alta", "A validar", "Confirmar politica Supabase/contratual.", ""),
    ("RNF-007", "Escalabilidade", "Preparado para crescimento de utilizadores e condominios.", "Media", "Implementado", "Rust/PostgreSQL; indices e checks a manter.", ""),
    ("RNF-008", "Mobile", "Uso em telemovel/PWA sem instalacao complexa.", "Alta", "Implementado", "PWA pronta para instalar.", ""),
    ("RNF-009", "Qualidade", "Smoke tests por API e por perfil de utilizador.", "Alta", "Implementado", "Contrato interno: API primeiro, depois HQ/Worker/Client.", ""),
    ("RNF-010", "Operacao", "Sem localhost como validacao final.", "Alta", "Implementado", "Validacao deve usar URL publicada/preview.", ""),
]


QUESTIONS = [
    "Que perfis existem na empresa e que decisoes cada perfil pode tomar?",
    "Quais os estados oficiais de um pedido, desde a entrada ate ao fecho?",
    "Que tempos de resposta/SLA querem medir por tipo de pedido?",
    "Que documentos sao obrigatorios por condominio, fraccao, fornecedor e pedido?",
    "Que relatorios a administracao precisa mensalmente e quais sao para clientes?",
    "Como funciona hoje a contabilidade: quotas, recibos, dividas, despesas e aprovacao?",
    "Que notificacoes fazem sentido: email, push, SMS, WhatsApp ou apenas dentro da app?",
    "Que informacao o cliente pode ver sem expor dados internos da equipa?",
    "Que integracoes sao obrigatorias para fechar a primeira fase?",
    "Que dados historicos precisam de ser importados antes do arranque?",
]


DEMO_STEPS = [
    ("1", "Abrir HQ", "Entrar em https://gestisac-web.vercel.app/hq/login e confirmar que a sessao de demonstracao abre sem escrever credenciais.", "A entrada sem friccao serve para demonstracao, desenvolvimento e smoke tests; producao real mantem controlo de acesso."),
    ("2", "Hoje / Dashboard", "Mostrar indicadores, prioridades e os quatro cartoes/resumos principais.", "Esta e a mesa de comando: a administracao ve o que precisa de atencao antes de entrar no detalhe."),
    ("3", "Condominios", "Abrir a lista, selecionar um condominio e mostrar o painel lateral.", "O condominio deixa de ser uma ficha isolada: liga pedidos, tarefas, agenda, documentos e equipa."),
    ("4", "Pedidos", "Selecionar um pedido, ver estado, responsavel e historico.", "O pedido passa a ter vida operacional: entra, e triado, atribuido, executado, validado e fica registado."),
    ("5", "Equipa", "Abrir a equipa e selecionar um membro.", "A administracao consegue perceber carga de trabalho, validacoes pendentes e distribuicao do dia."),
    ("6", "Tarefas", "Filtrar por Hoje/Em curso/Validacao/Atrasadas.", "Tarefas nao e uma tabela nova; e uma vista operacional sobre trabalho real ja existente."),
    ("7", "Agenda", "Mostrar proximos eventos e ligacoes.", "A agenda transforma tarefas e pedidos em planeamento visivel."),
    ("8", "Worker", "Abrir https://gestisac-web.vercel.app/worker/login.", "O funcionario ve o que tem de executar, sem acesso a areas administrativas indevidas."),
    ("9", "Client", "Abrir https://gestisac-web.vercel.app/client/login.", "O cliente acompanha o que lhe diz respeito, com linguagem simples e sem informacao interna."),
    ("10", "Fecho", "Abrir a matriz de requisitos.", "A reuniao fecha com validacao: isto esta feito, isto esta parcial, isto esta planeado, e isto precisa da vossa decisao."),
]


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)


def font(size: int = 24, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline="#CBD5E1", radius=20, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw: ImageDraw.ImageDraw, box, text, fill, fnt):
    lines = textwrap.wrap(text, width=24)
    line_h = fnt.getbbox("Ag")[3] + 8
    total = line_h * len(lines)
    y = box[1] + ((box[3] - box[1] - total) / 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = box[0] + ((box[2] - box[0] - (bbox[2] - bbox[0])) / 2)
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, fill="#2E74B5", width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    direction = 1 if x2 >= x1 else -1
    head = [(x2, y2), (x2 - 18 * direction, y2 - 10), (x2 - 18 * direction, y2 + 10)]
    draw.polygon(head, fill=fill)


def save_architecture_diagram() -> Path:
    path = ASSETS / "arquitetura-sistema.png"
    img = Image.new("RGB", (1600, 900), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title = font(48, True)
    body = font(28, False)
    small = font(22, False)
    draw.text((70, 55), "Arquitetura GESTISAC em producao", font=title, fill=f"#{DARK}")

    boxes = [
        ((90, 230, 360, 390), "Utilizadores\nHQ / Funcionario / Cliente", "#FFFFFF"),
        ((470, 230, 740, 390), "PWA / Frontend\nQwik na Vercel", "#E8EEF5"),
        ((850, 230, 1120, 390), "API\nRust + Axum", "#FFFFFF"),
        ((1230, 230, 1500, 390), "Base de dados\nPostgreSQL / Supabase", "#E8EEF5"),
    ]
    for box, label, fill in boxes:
        rounded_box(draw, box, fill, outline="#94A3B8")
        centered_text(draw, box, label, f"#{DARK}", body)
    draw_arrow(draw, (360, 310), (470, 310))
    draw_arrow(draw, (740, 310), (850, 310))
    draw_arrow(draw, (1120, 310), (1230, 310))

    bottom = [
        ((235, 590, 525, 720), "GitHub\ncontrolo de versoes", "#FFFFFF"),
        ((655, 590, 945, 720), "Checks e smoke tests\nAPI publicada + perfis", "#FFFFFF"),
        ((1075, 590, 1365, 720), "Deploy / warmup\nWeb + API", "#FFFFFF"),
    ]
    for box, label, fill in bottom:
        rounded_box(draw, box, fill, outline="#CBD5E1")
        centered_text(draw, box, label, f"#{DARK}", small)
    draw_arrow(draw, (525, 655), (655, 655), fill=f"#{GOLD}")
    draw_arrow(draw, (945, 655), (1075, 655), fill=f"#{GOLD}")
    draw.text((80, 805), "Leitura simples: a app corre no browser/PWA, comunica com a API Rust, e a API guarda/consulta dados estruturados em PostgreSQL.", font=small, fill="#475467")
    img.save(path)
    return path


def save_data_diagram() -> Path:
    path = ASSETS / "modelo-dados-simplificado.png"
    img = Image.new("RGB", (1600, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = font(46, True)
    body = font(24, False)
    small = font(20, False)
    draw.text((70, 55), "Modelo de dados simplificado", font=title, fill=f"#{DARK}")

    nodes = {
        "Tenant / Empresa": (640, 130, 960, 240),
        "Users / Equipa": (130, 360, 410, 470),
        "Condominios": (500, 360, 780, 470),
        "Fraccoes": (870, 360, 1150, 470),
        "Tickets / Pedidos": (1240, 360, 1520, 470),
        "Tarefas compostas": (130, 650, 410, 760),
        "Agenda": (500, 650, 780, 760),
        "Documentos": (870, 650, 1150, 760),
        "Contabilidade": (1240, 650, 1520, 760),
    }
    for name, box in nodes.items():
        fill = "#E8EEF5" if name in {"Tenant / Empresa", "Condominios", "Tickets / Pedidos"} else "#F8FAFC"
        rounded_box(draw, box, fill, outline="#94A3B8")
        centered_text(draw, box, name, f"#{DARK}", body)
    center = (800, 240)
    for target in ["Users / Equipa", "Condominios", "Tickets / Pedidos"]:
        b = nodes[target]
        draw_arrow(draw, center, ((b[0] + b[2]) // 2, b[1]), fill="#64748B", width=3)
    draw_arrow(draw, (780, 415), (870, 415), fill=f"#{MID_BLUE}", width=4)
    draw_arrow(draw, (1150, 415), (1240, 415), fill=f"#{MID_BLUE}", width=4)
    draw_arrow(draw, (410, 705), (500, 705), fill=f"#{GOLD}", width=4)
    draw_arrow(draw, (780, 705), (870, 705), fill=f"#{GOLD}", width=4)
    draw_arrow(draw, (1150, 705), (1240, 705), fill=f"#{GOLD}", width=4)
    draw.text((90, 855), "Nota: Tarefas e uma vista operacional composta. Nao exige uma tabela nova nesta fase.", font=small, fill="#475467")
    draw.text((90, 895), "O objetivo e manter a fonte de verdade em dados existentes: utilizadores, tickets, vistorias, manutencao e agenda.", font=small, fill="#475467")
    img.save(path)
    return path


def save_flow_diagram() -> Path:
    path = ASSETS / "fluxo-pedido-fim-a-fim.png"
    img = Image.new("RGB", (1600, 700), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title = font(44, True)
    body = font(23, False)
    small = font(19, False)
    draw.text((70, 50), "Fluxo de pedido do inicio ao fim", font=title, fill=f"#{DARK}")
    steps = [
        ("Cliente / HQ", "cria pedido"),
        ("Triagem", "classifica prioridade"),
        ("Atribuicao", "responsavel definido"),
        ("Execucao", "funcionario atualiza"),
        ("Validacao", "administracao confirma"),
        ("Historico", "registo auditavel"),
    ]
    x = 70
    y = 250
    w = 185
    h = 130
    for idx, (label, detail) in enumerate(steps):
        box = (x, y, x + w, y + h)
        rounded_box(draw, box, "#FFFFFF" if idx % 2 == 0 else "#E8EEF5", outline="#94A3B8")
        centered_text(draw, (box[0] + 10, box[1] + 18, box[2] - 10, box[1] + 74), label, f"#{DARK}", body)
        centered_text(draw, (box[0] + 10, box[1] + 78, box[2] - 10, box[3] - 12), detail, "#475467", small)
        if idx < len(steps) - 1:
            draw_arrow(draw, (x + w, y + h // 2), (x + w + 34, y + h // 2), fill=f"#{MID_BLUE}", width=4)
        x += w + 55
    draw.text((80, 570), "Mensagem para cliente: cada pedido deixa de depender de telefonemas soltos; passa a ter dono, estado, historico e proxima acao.", font=small, fill="#475467")
    img.save(path)
    return path


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, header_fill: str = LIGHT) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_fill(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(DARK)


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, MID_BLUE, 16, 8),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "GESTISAC | Dossier executivo-tecnico"
    header.style = doc.styles["Header"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "Documento de apoio a reuniao comercial e validacao de requisitos"
    footer.style = doc.styles["Footer"]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("GESTISAC")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Dossier formal + apresentacao executiva-tecnica")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    rows = [
        ("Objetivo", "Apoiar fecho de contrato e iniciar validacao formal de requisitos."),
        ("Publico", "Dona/decisora da empresa, equipa administrativa e interlocutores operacionais."),
        ("Nivel", "Executivo-tecnico: sem codigo, mas com arquitetura, fluxos, seguranca e roadmap."),
        ("Data", TODAY),
        ("Ambiente demonstravel", "https://gestisac-web.vercel.app"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.75)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, True, DARK, 10)
        set_cell_text(table.cell(idx, 1), value, False, "000000", 10)
        set_cell_fill(table.cell(idx, 0), SOFT_BLUE if idx % 2 == 0 else LIGHT)
    style_table(table, SOFT_BLUE)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "Mensagem central: o GESTISAC organiza o dia a dia da administracao de condominios num unico sistema, "
        "ligando condominio, equipa, tarefas, pedidos, agenda, documentos e dados de gestao."
    )
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_dossier(architecture: Path, data_model: Path, flow: Path) -> Path:
    doc = Document()
    set_doc_styles(doc)
    add_header_footer(doc)
    add_cover(doc)

    doc.add_heading("1. Visao geral", level=1)
    doc.add_paragraph(
        "O GESTISAC foi desenhado para reduzir a dispersao operacional de uma administracao de condominios. "
        "Em vez de pedidos em chamadas, mensagens soltas, folhas de calculo e documentos espalhados, a plataforma "
        "organiza a operacao em torno de seis janelas principais: Hoje, Condominios, Equipa, Tarefas, Pedidos e Agenda."
    )
    doc.add_heading("Objetivos do sistema", level=2)
    add_bullets(
        doc,
        [
            "Dar a administracao uma visao diaria clara do que exige atencao.",
            "Ligar pedidos, tarefas, equipa e agenda aos condominios reais.",
            "Separar a experiencia por perfil: HQ/Admin, funcionario e cliente.",
            "Criar base tecnica para documentos, contabilidade, relatorios e auditoria.",
            "Permitir demonstracao e validacao sem confundir app atual com roadmap futuro.",
        ],
    )
    doc.add_heading("Beneficios por perfil", level=2)
    table = doc.add_table(rows=1, cols=3)
    headers = ["Administracao / HQ", "Funcionarios", "Clientes"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, True, DARK)
    for row in [
        ("Acompanha prioridades, carga da equipa, pedidos e agenda.", "Ve tarefas atribuidas e atualiza execucao.", "Consulta pedidos e acompanha estado."),
        ("Reduz dependencia de mensagens soltas e memoria individual.", "Tem menos ambiguidade sobre a proxima acao.", "Ganha transparencia sem expor dados internos."),
        ("Fica com historico e base para relatorios.", "Regista atividade ligada ao pedido/condominio.", "Sabe quando o pedido avancou ou ficou pendente."),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False)
    style_table(table)

    doc.add_heading("App demonstravel vs. roadmap", level=2)
    doc.add_paragraph(
        "A app demonstravel ja permite apresentar o fluxo principal, a navegacao limpa e a separacao por perfis. "
        "O roadmap cobre modulos que exigem validacao das regras internas da empresa, especialmente contabilidade, "
        "relatorios, documentos avancados, notificacoes e integracoes."
    )

    doc.add_heading("2. Funcionalidades do sistema", level=1)
    table = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["Area", "Finalidade", "Utilizadores", "Estado", "Evolucao prevista"]):
        set_cell_text(table.cell(0, i), h, True, DARK, 8)
    for module in MODULES:
        cells = table.add_row().cells
        values = [module["area"], module["purpose"], module["users"], module["state"], module["next"]]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, i in {0, 3}, DARK if i in {0, 3} else "000000", 8)
    style_table(table)

    doc.add_heading("3. Fluxos de utilizacao", level=1)
    doc.add_picture(str(flow), width=Inches(6.35))
    doc.add_heading("Fluxo HQ / Administracao", level=2)
    add_numbered(
        doc,
        [
            "Entrar na app HQ sem credenciais manuais em demonstracao.",
            "Abrir Hoje para perceber prioridades e carga operacional.",
            "Selecionar Condominios para consultar ficha operacional.",
            "Abrir Pedidos para acompanhar estado, responsavel e historico.",
            "Usar Equipa/Tarefas para redistribuir trabalho e validar execucao.",
            "Consultar Agenda para planear visitas, prazos e reunioes.",
        ],
    )
    doc.add_heading("Fluxo Funcionario", level=2)
    add_numbered(
        doc,
        [
            "Entrar no contexto Worker.",
            "Ver tarefas e pedidos atribuidos.",
            "Atualizar estado, notas e execucao.",
            "Confirmar agenda e proximos compromissos.",
        ],
    )
    doc.add_heading("Fluxo Cliente", level=2)
    add_numbered(
        doc,
        [
            "Entrar no contexto Client.",
            "Consultar pedidos e estado.",
            "Acompanhar comunicacao relevante e proximos eventos quando aplicavel.",
        ],
    )

    doc.add_heading("4. Requisitos do sistema", level=1)
    doc.add_paragraph(
        "A matriz completa segue como ficheiro proprio para validacao em reuniao. Abaixo esta uma leitura resumida por tipo de requisito."
    )
    summary = {}
    for _, area, _, _, state, _, _ in REQUIREMENTS:
        summary.setdefault(area, {"Implementado": 0, "Parcial": 0, "Planeado": 0, "A validar": 0})
        summary[area][state] += 1
    table = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["Area", "Implementado", "Parcial", "Planeado", "A validar"]):
        set_cell_text(table.cell(0, i), h, True, DARK)
    for area, counts in summary.items():
        cells = table.add_row().cells
        set_cell_text(cells[0], area, True, DARK)
        for i, state in enumerate(["Implementado", "Parcial", "Planeado", "A validar"], start=1):
            set_cell_text(cells[i], str(counts[state]), False)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)

    doc.add_heading("Requisitos nao funcionais principais", level=2)
    add_bullets(
        doc,
        [
            "Seguranca: separacao por perfil, API protegida, sem segredos no frontend.",
            "Performance: frontend Qwik/PWA, API Rust e PostgreSQL com indices e checks.",
            "Disponibilidade: deploy publicado, smoke tests e warmup para reduzir arranque frio.",
            "Mobile/PWA: utilizacao em telemovel sem obrigar a app nativa na primeira fase.",
            "Auditoria: historico e rastreabilidade como base para validacao e relatorios.",
            "Backups: politica final a confirmar no contrato e configuracao Supabase.",
        ],
    )
    doc.add_heading("Perguntas para levantamento", level=2)
    add_bullets(doc, QUESTIONS)

    doc.add_heading("5. Arquitetura e tecnologia", level=1)
    doc.add_picture(str(architecture), width=Inches(6.35))
    doc.add_heading("Explicacao em linguagem simples", level=2)
    add_bullets(
        doc,
        [
            "Qwik/PWA: permite uma app rapida no browser e preparada para telemovel, com menos peso inicial.",
            "Rust/Axum: backend robusto, seguro e performante para regras de negocio e API.",
            "PostgreSQL/Supabase: base de dados relacional, estruturada e adequada a auditoria.",
            "Vercel/GitHub: deploy controlado, historico de versoes, checks automaticos e reposicao previsivel.",
        ],
    )
    doc.add_picture(str(data_model), width=Inches(6.35))
    doc.add_heading("Deployment e producao", level=2)
    doc.add_paragraph(
        "A versao publicada separa web, API e base de dados. Os checks de producao validam API publicada, readiness, "
        "smoke tests e fluxo por perfis. A demonstracao nao deve depender de localhost."
    )

    doc.add_heading("6. Estado atual e proximas fases", level=1)
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Feito / demonstravel", "Parcial / precisa validacao", "Planeado"]):
        set_cell_text(table.cell(0, i), h, True, DARK)
    rows = [
        (
            "Entrada sem credenciais manuais em demo; HQ/Worker/Client; menu limpo; Hoje, Equipa, Tarefas, Pedidos e Agenda.",
            "Ficha completa de condominios, documentos, validacoes de estados e regras de permissao finas.",
            "Contabilidade completa, relatorios executivos, notificacoes e integracoes externas.",
        ),
        (
            "API Rust publicada, PostgreSQL/Supabase, smoke tests e checks de producao.",
            "Auditoria detalhada, politicas de backup e criterios finais de SLA.",
            "Importacao historica, BI avancado, automacoes e processos financeiros finais.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, size=9)
    style_table(table)

    doc.add_heading("Riscos conhecidos e decisoes pendentes", level=2)
    add_bullets(
        doc,
        [
            "Nao prometer contabilidade final sem levantar regras reais da empresa.",
            "Definir matriz de permissoes antes de abrir documentos sensiveis a clientes.",
            "Validar modelos de relatorio e linguagem de comunicacao com residentes.",
            "Confirmar politicas de backup, retencao de documentos e responsabilidades contratuais.",
        ],
    )
    doc.add_heading("Proximas fases sugeridas", level=2)
    add_numbered(
        doc,
        [
            "Reuniao de validacao com decisora e equipa operacional.",
            "Fecho da matriz de requisitos obrigatorios para primeira fase contratual.",
            "Priorizacao de contabilidade, documentos e relatorios.",
            "Plano de importacao de dados e piloto com condominios reais.",
            "Go-live controlado com smoke tests e criterios de aceitacao assinados.",
        ],
    )

    docx_path = OUT / "GESTISAC_Dossier_Formal.docx"
    doc.save(docx_path)
    return docx_path


def build_matrix() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Matriz requisitos"
    headers = ["Codigo", "Area", "Requisito", "Prioridade", "Estado", "Observacoes", "Validacao Cliente"]
    ws.append(headers)
    for row in REQUIREMENTS:
        ws.append(row)

    header_fill = PatternFill("solid", fgColor=SOFT_BLUE)
    thin = Side(style="thin", color="CBD5E1")
    status_fills = {
        "Implementado": PatternFill("solid", fgColor="E7F4E8"),
        "Parcial": PatternFill("solid", fgColor="FFF3D6"),
        "Planeado": PatternFill("solid", fgColor="E8EEF5"),
        "A validar": PatternFill("solid", fgColor="FDEAEA"),
    }
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if cell.row == 1:
                cell.fill = header_fill
                cell.font = Font(bold=True, color=DARK)
            elif cell.column == 5:
                cell.fill = status_fills.get(str(cell.value), PatternFill())
                cell.font = Font(bold=True)
    widths = [12, 18, 58, 14, 16, 46, 26]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    q = wb.create_sheet("Perguntas validacao")
    q.append(["Tema", "Pergunta", "Resposta / decisao cliente"])
    for item in QUESTIONS:
        theme = item.split()[1] if " " in item else "Geral"
        q.append(["Validacao", item, ""])
    for row in q.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if cell.row == 1:
                cell.fill = header_fill
                cell.font = Font(bold=True, color=DARK)
    q.column_dimensions["A"].width = 18
    q.column_dimensions["B"].width = 78
    q.column_dimensions["C"].width = 46
    q.freeze_panes = "A2"

    path = OUT / "GESTISAC_Matriz_Requisitos.xlsx"
    wb.save(path)
    return path


def build_demo_script() -> Path:
    path = OUT / "GESTISAC_Roteiro_Demo.md"
    lines = [
        "# GESTISAC - Roteiro de demo",
        "",
        f"Data de preparacao: {TODAY}",
        "",
        "## Objetivo da demo",
        "",
        "Mostrar que o GESTISAC ja organiza a operacao diaria em seis janelas principais, separa perfis de utilizador e cria base para validar requisitos de contrato.",
        "",
        "## Ordem recomendada",
        "",
    ]
    for number, title, action, message in DEMO_STEPS:
        lines.extend(
            [
                f"### {number}. {title}",
                "",
                f"- Clique/acao: {action}",
                f"- Mensagem-chave: {message}",
                "",
            ]
        )
    lines.extend(
        [
            "## Frases de apoio",
            "",
            "- \"Esta versao serve para ver a operacao real e validar regras antes de fechar tudo em contrato.\"",
            "- \"O que esta marcado como implementado e demonstravel; o que esta parcial ou planeado nao deve ser vendido como fechado.\"",
            "- \"A matriz de requisitos e a ferramenta para a cliente dizer: isto sim, isto nao, isto falta.\"",
            "",
            "## Links",
            "",
            "- HQ: https://gestisac-web.vercel.app/hq/login",
            "- Worker: https://gestisac-web.vercel.app/worker/login",
            "- Client: https://gestisac-web.vercel.app/client/login",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def build_readme(paths: dict[str, Path]) -> Path:
    path = OUT / "README.md"
    lines = [
        "# Pacote de apresentacao GESTISAC",
        "",
        "Entregaveis preparados para reuniao comercial e validacao formal de requisitos.",
        "",
        "## Ficheiros",
        "",
    ]
    for label, file_path in paths.items():
        lines.append(f"- {label}: `{file_path.name}`")
    lines.extend(
        [
            "",
            "## Nota de uso",
            "",
            "Usar o dossier como documento para deixar com a cliente. Usar os slides para conduzir a reuniao. Usar a matriz para recolher decisoes e fechar requisitos.",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def status_class(status: str) -> str:
    return {
        "Implementado": "done",
        "Parcial": "partial",
        "Planeado": "planned",
        "A validar": "validate",
    }.get(status, "validate")


def build_html(architecture: Path, data_model: Path, flow: Path) -> Path:
    path = OUT / "GESTISAC_Prototipo_Apresentacao.html"
    module_cards = "\n".join(
        f"""
        <article class="module-card">
          <header><span>{escape(module['area'])}</span><strong class="{status_class(module['state'])}">{escape(module['state'])}</strong></header>
          <p>{escape(module['purpose'])}</p>
          <small>{escape(module['next'])}</small>
        </article>
        """
        for module in MODULES
    )
    req_rows = "\n".join(
        f"""
        <tr>
          <td>{escape(code)}</td>
          <td>{escape(area)}</td>
          <td>{escape(req)}</td>
          <td>{escape(priority)}</td>
          <td><span class="badge {status_class(state)}">{escape(state)}</span></td>
          <td>{escape(obs)}</td>
          <td class="empty-cell">{escape(validation)}</td>
        </tr>
        """
        for code, area, req, priority, state, obs, validation in REQUIREMENTS
    )
    demo_rows = "\n".join(
        f"""
        <li>
          <b>{escape(number)}. {escape(title)}</b>
          <span>{escape(action)}</span>
          <em>{escape(message)}</em>
        </li>
        """
        for number, title, action, message in DEMO_STEPS
    )
    questions = "\n".join(f"<li>{escape(q)}</li>" for q in QUESTIONS)
    html = f"""<!doctype html>
<html lang="pt">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>GESTISAC - Prototipo de apresentacao</title>
  <style>
    :root {{
      --ink: #0b2545;
      --navy: #163b66;
      --blue: #2e74b5;
      --sky: #e8eef5;
      --paper: #f8fafc;
      --line: #cbd5e1;
      --muted: #667085;
      --gold: #d7a84f;
      --green: #2e7d32;
      --amber: #b36b00;
      --red: #9b1c1c;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: "Segoe UI", Arial, sans-serif;
      color: var(--ink);
      background: var(--paper);
      letter-spacing: 0;
    }}
    a {{ color: inherit; }}
    .topbar {{
      position: sticky;
      top: 0;
      z-index: 10;
      display: flex;
      align-items: center;
      gap: 18px;
      padding: 14px 28px;
      border-bottom: 1px solid var(--line);
      background: rgba(248, 250, 252, 0.94);
      backdrop-filter: blur(12px);
    }}
    .brand {{ font-weight: 800; letter-spacing: .02em; }}
    .topbar nav {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-left: auto;
    }}
    .topbar a {{
      text-decoration: none;
      font-size: 13px;
      font-weight: 650;
      padding: 8px 10px;
      border: 1px solid var(--line);
      background: white;
    }}
    .section {{
      min-height: calc(100vh - 58px);
      padding: 58px clamp(28px, 6vw, 92px);
      border-bottom: 1px solid var(--line);
    }}
    .kicker {{
      color: var(--gold);
      text-transform: uppercase;
      font-weight: 800;
      font-size: 13px;
      margin-bottom: 14px;
    }}
    h1, h2 {{
      margin: 0;
      color: var(--ink);
      line-height: 1.04;
      max-width: 960px;
    }}
    h1 {{ font-size: clamp(44px, 7vw, 84px); }}
    h2 {{ font-size: clamp(34px, 5vw, 58px); }}
    .lead {{
      max-width: 880px;
      color: var(--muted);
      font-size: clamp(18px, 2.2vw, 24px);
      line-height: 1.42;
      margin: 18px 0 0;
    }}
    .metric-row, .lane-grid, .profile-grid, .status-grid {{
      display: grid;
      gap: 18px;
      margin-top: 48px;
    }}
    .metric-row {{ grid-template-columns: repeat(3, minmax(0, 1fr)); }}
    .lane-grid {{ grid-template-columns: repeat(4, minmax(0, 1fr)); }}
    .profile-grid, .status-grid {{ grid-template-columns: repeat(3, minmax(0, 1fr)); }}
    .metric, .lane, .profile, .status-card, .module-card {{
      background: white;
      border: 1px solid var(--line);
      padding: 22px;
      min-height: 150px;
    }}
    .metric {{ border-left: 7px solid var(--blue); }}
    .metric:nth-child(2) {{ border-left-color: var(--gold); }}
    .metric:nth-child(3) {{ border-left-color: var(--green); }}
    .metric b {{
      display: block;
      font-size: 32px;
      margin-bottom: 10px;
    }}
    .metric span, .lane p, .profile p, .status-card p, .module-card p, .module-card small {{
      color: var(--muted);
      line-height: 1.45;
    }}
    .lane {{ border-top: 7px solid var(--blue); }}
    .lane:nth-child(1) {{ border-top-color: var(--red); }}
    .lane:nth-child(2) {{ border-top-color: var(--amber); }}
    .lane:nth-child(4) {{ border-top-color: var(--green); }}
    .flow {{
      display: grid;
      grid-template-columns: repeat(6, 1fr);
      gap: 12px;
      margin-top: 52px;
    }}
    .flow div {{
      background: white;
      border: 1px solid var(--line);
      min-height: 140px;
      padding: 20px;
      position: relative;
    }}
    .flow b {{
      display: block;
      color: var(--gold);
      font-size: 28px;
      margin-bottom: 18px;
    }}
    .demo-table, .requirements {{
      width: 100%;
      border-collapse: collapse;
      margin-top: 42px;
      background: white;
      font-size: 14px;
    }}
    .demo-table th, .demo-table td, .requirements th, .requirements td {{
      border: 1px solid var(--line);
      padding: 13px 14px;
      text-align: left;
      vertical-align: top;
    }}
    .demo-table th, .requirements th {{
      background: var(--ink);
      color: white;
    }}
    .badge, .done, .partial, .planned, .validate {{
      display: inline-block;
      font-weight: 800;
      font-size: 12px;
    }}
    .done {{ color: var(--green); }}
    .partial {{ color: var(--amber); }}
    .planned {{ color: var(--blue); }}
    .validate {{ color: var(--red); }}
    .badge {{
      padding: 5px 8px;
      background: var(--paper);
      border: 1px solid var(--line);
      white-space: nowrap;
    }}
    .diagram {{
      display: block;
      width: min(1040px, 100%);
      margin: 42px auto 0;
      border: 1px solid var(--line);
      background: white;
    }}
    .modules {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 14px;
      margin-top: 36px;
    }}
    .module-card {{ min-height: 170px; }}
    .module-card header {{
      display: flex;
      justify-content: space-between;
      gap: 16px;
      font-weight: 800;
      margin-bottom: 12px;
    }}
    .demo-script {{
      display: grid;
      gap: 10px;
      margin-top: 34px;
      padding: 0;
      list-style: none;
    }}
    .demo-script li {{
      display: grid;
      grid-template-columns: 230px 1fr 1.3fr;
      gap: 18px;
      padding: 15px 18px;
      border: 1px solid var(--line);
      background: white;
    }}
    .demo-script span {{ color: var(--muted); }}
    .demo-script em {{ color: var(--navy); font-style: normal; }}
    .questions {{
      columns: 2;
      column-gap: 48px;
      margin-top: 34px;
      font-size: 18px;
      line-height: 1.5;
    }}
    .questions li {{ margin-bottom: 12px; break-inside: avoid; }}
    .empty-cell {{ min-width: 170px; }}
    .note {{
      margin-top: 28px;
      color: var(--muted);
      font-size: 15px;
    }}
    @media (max-width: 900px) {{
      .metric-row, .lane-grid, .profile-grid, .status-grid, .modules, .flow {{
        grid-template-columns: 1fr;
      }}
      .demo-script li {{ grid-template-columns: 1fr; }}
      .questions {{ columns: 1; }}
      .topbar {{ align-items: flex-start; flex-direction: column; }}
      .topbar nav {{ margin-left: 0; }}
    }}
  </style>
</head>
<body>
  <header class="topbar">
    <div class="brand">GESTISAC</div>
    <nav>
      <a href="#abertura">Abertura</a>
      <a href="#operacao">Operacao</a>
      <a href="#demo">Demo</a>
      <a href="#arquitetura">Arquitetura</a>
      <a href="#requisitos">Requisitos</a>
      <a href="#roteiro">Roteiro</a>
    </nav>
  </header>

  <section id="abertura" class="section">
    <div class="kicker">Prototipo visual</div>
    <h1>GESTISAC</h1>
    <p class="lead">Operacao de condominios organizada, visivel e validavel. Esta pagina serve para rever a narrativa antes de decidir se a apresentacao final fica em HTML ou slides.</p>
    <div class="metric-row">
      <article class="metric"><b>6</b><span>Janelas principais: Hoje, Condominios, Equipa, Tarefas, Pedidos e Agenda.</span></article>
      <article class="metric"><b>3</b><span>Perfis separados: HQ/Admin, Funcionario e Cliente.</span></article>
      <article class="metric"><b>Qwik + Rust + PostgreSQL</b><span>PWA, API robusta e dados estruturados.</span></article>
    </div>
  </section>

  <section class="section">
    <div class="kicker">Problema</div>
    <h2>A operacao perde tempo quando a informacao vive fora do sistema</h2>
    <p class="lead">A conversa com a cliente deve comecar pelo custo operacional: pedidos dispersos, falta de visibilidade, dados desligados e clientes a perguntar pelo estado.</p>
    <div class="lane-grid">
      <article class="lane"><h3>Pedidos dispersos</h3><p>Chamadas, emails e mensagens dificultam dono, prioridade e historico.</p></article>
      <article class="lane"><h3>Equipa sem visibilidade</h3><p>A administracao nao ve facilmente carga, validacoes e atrasos.</p></article>
      <article class="lane"><h3>Dados desconectados</h3><p>Condominio, agenda, documentos, tarefas e contabilidade ficam em silos.</p></article>
      <article class="lane"><h3>Cliente sem estado</h3><p>Sem transparencia controlada, aumenta o trabalho manual de resposta.</p></article>
    </div>
  </section>

  <section id="operacao" class="section">
    <div class="kicker">Organizacao</div>
    <h2>O produto organiza a empresa por trabalho real</h2>
    <p class="lead">As seis janelas principais formam a camada diaria. Os modulos administrativos continuam ligados por contexto, sem dominar o menu.</p>
    <div class="modules">{module_cards}</div>
  </section>

  <section class="section">
    <div class="kicker">Fluxo diario</div>
    <h2>Da prioridade do dia ao historico auditavel</h2>
    <p class="lead">O valor esta em ligar cada acao a um responsavel, a um estado e a um contexto.</p>
    <div class="flow">
      <div><b>1</b><strong>Entrar</strong><p>sessao demo sem escrita manual</p></div>
      <div><b>2</b><strong>Priorizar</strong><p>Hoje mostra o que exige atencao</p></div>
      <div><b>3</b><strong>Atribuir</strong><p>pedido/tarefa ganha responsavel</p></div>
      <div><b>4</b><strong>Executar</strong><p>funcionario atualiza progresso</p></div>
      <div><b>5</b><strong>Validar</strong><p>administracao confirma resultado</p></div>
      <div><b>6</b><strong>Registar</strong><p>historico fica consultavel</p></div>
    </div>
    <img class="diagram" src="assets/fluxo-pedido-fim-a-fim.png" alt="Fluxo de pedido do inicio ao fim">
  </section>

  <section id="demo" class="section">
    <div class="kicker">Demo</div>
    <h2>As 6 janelas principais mostram a operacao completa</h2>
    <p class="lead">A demo deve seguir uma ordem simples: primeiro a administracao, depois equipa/trabalho, depois cliente.</p>
    <table class="demo-table">
      <thead><tr><th>Janela</th><th>Mensagem</th><th>Estado</th></tr></thead>
      <tbody>
        <tr><td>Hoje</td><td>Centro de comando e prioridades.</td><td><span class="done">Implementado</span></td></tr>
        <tr><td>Condominios</td><td>Ficha operacional e relacoes.</td><td><span class="partial">Parcial</span></td></tr>
        <tr><td>Equipa</td><td>Carga e atribuicoes.</td><td><span class="done">Implementado</span></td></tr>
        <tr><td>Tarefas</td><td>Vista composta do trabalho real.</td><td><span class="done">Implementado</span></td></tr>
        <tr><td>Pedidos</td><td>Ciclo de acompanhamento.</td><td><span class="done">Implementado</span></td></tr>
        <tr><td>Agenda</td><td>Planeamento e eventos.</td><td><span class="done">Implementado</span></td></tr>
      </tbody>
    </table>
  </section>

  <section id="arquitetura" class="section">
    <div class="kicker">Arquitetura</div>
    <h2>Tecnologia escolhida para rapidez, seguranca e evolucao</h2>
    <p class="lead">Explicacao simples para decisora: app no browser/PWA, API Rust, PostgreSQL/Supabase, Vercel/GitHub para deploy e validacao.</p>
    <img class="diagram" src="assets/arquitetura-sistema.png" alt="Arquitetura GESTISAC">
    <img class="diagram" src="assets/modelo-dados-simplificado.png" alt="Modelo de dados simplificado">
  </section>

  <section class="section">
    <div class="kicker">Estado atual</div>
    <h2>Separar feito, parcial e planeado evita promessas vagas</h2>
    <div class="status-grid">
      <article class="status-card"><h3 class="done">Feito / demonstravel</h3><p>Loginless demo, HQ/Worker/Client, menu limpo, Hoje, Equipa, Tarefas, Pedidos, Agenda, API publicada e checks.</p></article>
      <article class="status-card"><h3 class="partial">Parcial / validar</h3><p>Ficha profunda de condominios, documentos, estados finais, auditoria detalhada e permissoes finas.</p></article>
      <article class="status-card"><h3 class="planned">Planeado</h3><p>Contabilidade completa, relatorios executivos, notificacoes, integracoes e importacao historica.</p></article>
    </div>
  </section>

  <section id="requisitos" class="section">
    <div class="kicker">Matriz de requisitos</div>
    <h2>A matriz transforma a reuniao numa validacao concreta</h2>
    <p class="lead">Nao e para prometer tudo. E para marcar o que fica aprovado, o que muda e o que falta.</p>
    <table class="requirements">
      <thead><tr><th>Codigo</th><th>Area</th><th>Requisito</th><th>Prioridade</th><th>Estado</th><th>Observacoes</th><th>Validacao cliente</th></tr></thead>
      <tbody>{req_rows}</tbody>
    </table>
  </section>

  <section class="section">
    <div class="kicker">Perguntas</div>
    <h2>As melhores perguntas desbloqueiam o contrato</h2>
    <ol class="questions">{questions}</ol>
  </section>

  <section id="roteiro" class="section">
    <div class="kicker">Roteiro de demo</div>
    <h2>Ordem de cliques e mensagens-chave</h2>
    <ol class="demo-script">{demo_rows}</ol>
    <p class="note">Links reais: HQ https://gestisac-web.vercel.app/hq/login | Worker https://gestisac-web.vercel.app/worker/login | Client https://gestisac-web.vercel.app/client/login</p>
  </section>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")
    return path


def main() -> None:
    ensure_dirs()
    architecture = save_architecture_diagram()
    data_model = save_data_diagram()
    flow = save_flow_diagram()
    docx = build_dossier(architecture, data_model, flow)
    html = build_html(architecture, data_model, flow)
    matrix = build_matrix()
    demo = build_demo_script()
    readme = build_readme(
        {
            "Dossier DOCX": docx,
            "Prototipo visual HTML": html,
            "Matriz de requisitos": matrix,
            "Roteiro de demo": demo,
        }
    )
    for item in [docx, html, matrix, demo, readme, architecture, data_model, flow]:
        print(item)


if __name__ == "__main__":
    main()
